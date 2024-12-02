import requests
import urllib3
urllib3.disable_warnings()
import os
import base64
import json
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

api_username_token = os.environ.get("LE_GATEWAY_VLDN_USERNAME_TOKEN")
api_password_token = ""
host_ip = os.environ.get("LE_GATEWAY_VLDN_HOST_IP_ADDRESS")

# Generating the basic authorization token after taking LE API token as username and pwd values
def generate_basic_auth_token(username, password):
    credentials = f"{username}:{password}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()
    auth_token = f"Basic {encoded_credentials}"
    return auth_token

url = f"https://{host_ip}/dm/template"

###      dev and tags template upload       ###
with open("devices_and_tags_template.json", "r") as openfile:
    file = json.load(openfile)

payload = json.dumps(file)

headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

# Sending a PUT request with the JSON payload
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
print(response.text)

if response.status_code == 200 or response.status_code == 204:
    print("10 Devices and 500 tags each template uploaded successfully.")
else:
    print(f"Failed. Status code:", response.status_code)

time.sleep(10)

## enabling data logging ##

## Defining the list of device IDs
dev_Ids = [
    "247AA3B1-2719-466D-BD9E-72250D6841AA",
    "BC6028FA-4B28-4EC3-804A-3C1F5FD51DF2",
    "86FAC398-396A-458D-84B4-CDA0E0D7B877",
    "0E3D3F24-09E7-4757-98A8-F8B7E52B24EB",
    "7637AF1E-23C1-4032-BF9E-757C30268340",
    "84866870-B777-4BFD-A2A9-BA4CBD1D7CDC",
    "C7ED5DA1-8DB7-461E-A4CD-6E779E776A66",
    "1815AAC4-CBDD-4C64-8B0E-7DE42ABEDD58",
    "8A6FAC54-8AB5-4B02-86D2-9BD2B0376693",
    "D8483662-DD34-4ACC-A863-DE8D2826DD28",
]

# Defining common headers
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

# Iterating through device IDs and sending PUT requests with dynamic payloads
data_enab_count = 0
for i, dev_id in enumerate(dev_Ids, start=1):
    url = f"https://{host_ip}/stats/devices/{dev_id}"

    # Creating a dynamic payload with the "name" field
    payload = json.dumps(
        {
            "ID": dev_id,
            "name": f"test{i}",  # e.g., test1, test2, ...
            "retentionHours": 168,
            "saveAll": True,
        }
    )

    response = requests.put(url, headers=headers, data=payload, verify=False)
    if response.status_code == 204:
        data_enab_count += 1
    # print(f"Device ID: {dev_id}\nResponse: {response.text}\n")
if data_enab_count == 10:  # Data store enable success message
    print("Data store enable success")

time.sleep(10)


### analytics ###

url = f"https://{host_ip}/analytics/v2/import"

payload = json.dumps(
    {
        "Processors": [
            {
                "ID": "0290D0ED-BA51-415C-8033-47E0E2AE7C74",
                "Name": "Flow  - [6913]",
                "FunctionName": "Moving Maximum",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "field_name": "value",
                        "pass_through_value": False,
                        "processing_type": "",
                        "window_size": 10,
                    }
                },
                "InputEvents": ["CF333D5E-68A4-4A87-A335-EE08A4360311"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["1A180B0C-D7E3-4377-B21F-F641E88F37F2"],
            },
            {
                "ID": "094D8862-3C86-402D-8F76-6B70D7BF85E4",
                "Name": "Flow  - [3741]",
                "FunctionName": "Asset Online Percentage",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "device_reset": True,
                        "end_hour": 24,
                        "pass_through_value": False,
                        "start_hour": 0,
                        "timerInterval": 0,
                    }
                },
                "InputEvents": ["44180CA0-CF5A-4A5E-A210-72AFC5D03BCE"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["6940798B-7DAE-4E0C-BD65-87D1B5B7C652"],
            },
            {
                "ID": "164CDBFC-852C-493E-83F2-308BCFF56CD6",
                "Name": "Flow  - [1949]",
                "FunctionName": "DataHub Subscribe",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "format": "json",
                        "ignore_failed_data": False,
                        "ignore_null_value": False,
                        "topic": "devicehub.alias.test8.tag262",
                    }
                },
                "InputEvents": [],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["73256AB0-7236-4185-9518-F6C479AFA312"],
            },
            {
                "ID": "1A180B0C-D7E3-4377-B21F-F641E88F37F2",
                "Name": "Flow  - [6913]",
                "FunctionName": "DataHub Publish",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "single_topic": True,
                        "topic": "analytics.publish.sxZN5H6YEAwjDyoeRuroy",
                    }
                },
                "InputEvents": ["0290D0ED-BA51-415C-8033-47E0E2AE7C74"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": [],
            },
            {
                "ID": "44180CA0-CF5A-4A5E-A210-72AFC5D03BCE",
                "Name": "Flow  - [3741]",
                "FunctionName": "DataHub Subscribe",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "format": "json",
                        "ignore_failed_data": False,
                        "ignore_null_value": False,
                        "topic": "devicehub.alias.test4.tag455",
                    }
                },
                "InputEvents": [],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["094D8862-3C86-402D-8F76-6B70D7BF85E4"],
            },
            {
                "ID": "6940798B-7DAE-4E0C-BD65-87D1B5B7C652",
                "Name": "Flow  - [3741]",
                "FunctionName": "DataHub Publish",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "single_topic": True,
                        "topic": "analytics.publish.NEn_-Q78L0aE4H1WbG8mO",
                    }
                },
                "InputEvents": ["094D8862-3C86-402D-8F76-6B70D7BF85E4"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": [],
            },
            {
                "ID": "73256AB0-7236-4185-9518-F6C479AFA312",
                "Name": "Flow  - [1949]",
                "FunctionName": "Change of Value",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "delta_of_tolerance": 0,
                        "map_field_name": "value",
                        "pass_through_value": False,
                        "timerInterval": 0,
                    }
                },
                "InputEvents": ["164CDBFC-852C-493E-83F2-308BCFF56CD6"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["F91CFD41-B950-4344-8EB3-57768CE61D52"],
            },
            {
                "ID": "CF333D5E-68A4-4A87-A335-EE08A4360311",
                "Name": "Flow  - [6913]",
                "FunctionName": "DataHub Subscribe",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "format": "json",
                        "ignore_failed_data": False,
                        "ignore_null_value": False,
                        "topic": "devicehub.alias.test1.tag190",
                    }
                },
                "InputEvents": [],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": ["0290D0ED-BA51-415C-8033-47E0E2AE7C74"],
            },
            {
                "ID": "F91CFD41-B950-4344-8EB3-57768CE61D52",
                "Name": "Flow  - [1949]",
                "FunctionName": "DataHub Publish",
                "IsActive": True,
                "Config": {
                    "settings": {
                        "single_topic": True,
                        "topic": "analytics.publish.AT8Apo_nmiJmpcezkm06x",
                    }
                },
                "InputEvents": ["73256AB0-7236-4185-9518-F6C479AFA312"],
                "InputWaits": [],
                "InputsDefinitions": {},
                "State": None,
                "Outputs": [],
            },
        ],
        "Positions": [],
        "Version": "1.18.5",
        "Git": "8ff43348b2",
    }
)

headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

#response = requests.request("POST", url, headers=headers, data=payload, files=files,verify=False)

#headers = {
#    "Content-Type": "application/json",
#    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
#}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
if response.status_code == 204:
    print("Analytics template uploaded")

time.sleep(10)

###     flows   ####

url = f"https://{host_ip}/flows-manager/flows"

payload = json.dumps(
    {"id": None, "name": "gateway_validation", "secret": "", "memory_limit": 1024}
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)
#print(response)
resp1 = response.status_code
#print(resp1)

time.sleep(5)

url = f"https://{host_ip}/flows-manager/flows/2/start"

payload = json.dumps(
    {"id": 2, "name": "gateway_validation", "memory_limit": 1024, "secret": ""}
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
#print(response)
resp2 = response.status_code
#print(resp2)

if resp1 == 200 and resp2 == 204:
    print("Flows has been setup and is running")

time.sleep(10)

##  adding flows template   ##

url = f"https://{host_ip}/flows-2/flows"       #assuming no other flow is running

payload = {}
files=[
  ('UploadFile',('flows_template.json',open('flows_template.json','rb'),'application/json'))
]
headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("POST", url, headers=headers, data=payload, files=files,verify=False)
#print(response)
if(response.status_code==200):
    print("Flows template uploaded successfully")

time.sleep(15)
###     adding integrations  ###
# starting mysql from marketplace

url = f"https://{host_ip}/apps/marketplaces/default-marketplace/apps/090b4fc6-b95e-46c9-a7d7-78a0e7e8178e/versions/latest?locale=en-US"

payload = json.dumps(
    {
        "stack_name": "mysql",
        "stack_description": "",
        "MYSQL_PORT": "3306",
        "MYSQL_DATABASE": "litmus",
        "MYSQL_USER": "user",
        "MYSQL_PASSWORD": "loopedge-s3cr3t!",
        "MYSQL_ROOT_PASSWORD": "loopedge-s3cr3t!",
        "RESTART": "always",
    }
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)
if response.status_code == 200:
    print("Mysql marketplace container launched")
# print(response.text)
time.sleep(10)
###       adding mysql integration      ##
url = f"https://{host_ip}/cc/instances"

payload = json.dumps(
    {
        "providerId": "mysql",
        "config": '{"createTable":true,"database":"litmus","hostname":"{host_ip}","name":"mysql-integration","password":"loopedge-s3cr3t!","persistentStorage":true,"port":3306,"queueMode":"lifo","table":"new_table","throttlingLimit":0,"username":"root"}',
    }
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)

# print(response.text)
if response.status_code == 200:
    print("Mysql integration generated")
r_json = response.json()
if "instanceId" in r_json:
    extn = r_json["instanceId"]
# print(extn)
#       mysql connect enable

url = f"https://{host_ip}/cc/instances/{extn}/enable"
# print(url)
payload = {}
headers = {
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
if response.status_code == 204:
    print("Mysql integration enabled")

# print(response.text)
time.sleep(10)
#       adding mqtt integration

url = f"https://{host_ip}/cc/instances"

payload = json.dumps(
    {
        "providerId": "generic-mqtt-tcp",
        "config": '{"clientId":"gateway_vldn","hostname":"broker.hivemq.com","name":"hivemq","persistentStorage":true,"port":1883,"qos":1,"queueMode":"lifo","throttlingLimit":0,"topic":"litmus/topic/gateway","willPayloadType":"string","willQoS":1,"willRetained":false,"workersCount":100}',
    }
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)

# print(response.text)
r_json = response.json()
if "instanceId" in r_json:
    extn = r_json["instanceId"]
# print(extn)
if response.status_code == 200:
    print("Mqtt integration generated")

# Use .format() to insert extn into the URL
url = f"https://{host_ip}/cc/instances/{extn}/enable"
# print(url)

payload = {}
headers = {
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
if response.status_code == 204:
    print("Mqtt integration enabled")

time.sleep(15)

###         sending topic to cloud      ####

#get instance ids

url = f"https://{host_ip}/cc/instances"

payload = {}
headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("GET", url, headers=headers, data=payload, verify=False)
#print('get instances response:',response)

resp_json=response.json()

# Initialize variables
inst_mysql = None
inst_mqtt = None

# Iterate through the JSON data to find the instances with 'providerId' being 'mysql' and 'generic-mqtt-tcp'
for data in resp_json:
    if data['providerId'] == 'mysql':
        inst_mysql = data['instanceId']
    elif data['providerId'] == 'generic-mqtt-tcp':
        inst_mqtt = data['instanceId']


##  import topics for mysql     ##

#disable first
url = f"https://{host_ip}/cc/instances/{inst_mysql}/disable"

payload={}
headers={
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)

#iterative imports
url = f"https://{host_ip}/cc/instances/{inst_mysql}/import"

#print(url)
ctr_mysql=0
for i in range(10):
    payload = json.dumps({
    "deviceId": dev_Ids[i],
    "wildcard": False
    })
    headers = {
    'Content-Type': 'application/json',
    'Authorization': generate_basic_auth_token(api_username_token, api_password_token),
    }
    response = requests.request("POST", url, headers=headers, data=payload, verify=False)
    if(response.status_code==200):
        ctr_mysql+=1
if(ctr_mysql==10):
    print('Topics imported successfully for Mysql integration')

time.sleep(10)

#enable and get topics ids

url = f"https://{host_ip}/cc/instances/{inst_mysql}/enable"

payload={}
headers={
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)

url = "https://{host_ip}/cc/instances/{}/subs".format(inst_mysql)

payload = {}
headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

# Make the GET request
response = requests.get(url, headers=headers, data=payload, verify=False)

# Check if the request was successful
if response.status_code == 200:
    # Parse the JSON response
    res_topics = response.json()

    # Initialize an empty list to store 'topicsId' values
    topics_list_mysql = []

    # Iterate through the response and extract 'topicsId' values
    for item in res_topics:
        if 'topicId' in item:
            topics_list_mysql.append(item['topicId'])

#       enable to cloud for mysql
#disable first
url = f"https://{host_ip}/cc/instances/{inst_mysql}/disable"

payload={}
headers={
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)

url = f"https://{host_ip}/cc/instances/{inst_mysql}/enableSubs"

payload=json.dumps(topics_list_mysql)
headers={
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
#print(response)
if(response.status_code==200):
    print("Topics enabled to cloud for mysql successfully")

time.sleep(10)
#enable final for mysql
url = f"https://{host_ip}/cc/instances/{inst_mysql}/enable"

payload={}
headers={
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}
response = requests.request("PUT", url, headers=headers, data=payload, verify=False)

time.sleep(10)
##      import topic for mqtt  ##

url = f"https://{host_ip}/cc/instances/{inst_mqtt}/import"
#print(url)

ctr=0
for i in range(10):
    #print(dev_Ids[i])
    payload = json.dumps({
    "deviceId": dev_Ids[i],
    "wildcard": False
    })
    headers = {
    'Content-Type': 'application/json',
    'Authorization': generate_basic_auth_token(api_username_token, api_password_token),
    }

    response = requests.request("POST", url, headers=headers, data=payload,verify=False)
    if(response.status_code==200):
        ctr+=1
if(ctr==10):
    print("Topics imported successfully for Mqtt integration")
time.sleep(10)

url = f"https://{host_ip}/cc/instances/{inst_mqtt}/subs"

payload = {}
headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

# Make the GET request
response = requests.get(url, headers=headers, data=payload, verify=False)

# Check if the request was successful
if response.status_code == 200:
    # Parse the JSON response
    res_topics = response.json()

    # Initialize an empty list to store 'topicsId' values
    topics_list = []

    # Iterate through the response and extract 'topicsId' values
    for item in res_topics:
        if 'topicId' in item:
            topics_list.append(item['topicId'])

#print('Mqtt topics',topics_list)
##      enable those topics
url = f"https://{host_ip}/cc/instances/{inst_mqtt}/enableSubs"
payload = json.dumps(topics_list)
 
headers = {
  'Authorization': generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("PUT", url, headers=headers, data=payload,verify=False)
#print(response)
if(response.status_code==200):
    print("Topics enabled to cloud for mqtt successfully")

time.sleep(10)

###         opc ua      ###

# setup management settings

url = f"https://{host_ip}/opcua/modes"

payload = json.dumps(
    [
        {"name": "None-None", "enabled": True},
        {"name": "Basic256Sha256-SignAndEncrypt", "enabled": False},
        {"name": "Basic256Sha256-Sign", "enabled": False},
        {"name": "Basic256-SignAndEncrypt", "enabled": False},
        {"name": "Basic256-Sign", "enabled": False},
        {"name": "Basic128Rsa15-SignAndEncrypt", "enabled": False},
        {"name": "Basic128Rsa15-Sign", "enabled": False},
    ]
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)

#print(response)

url = f"https://{host_ip}/opcua/policies"

payload = json.dumps(
    [{"Authentication": "Anonymous", "EnabledPolicies": [{"name": "None"}]}]
)
headers = {
    "Content-Type": "application/json",
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token),
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
#print(response)
#        import dh tags onto hierarchy      ##
time.sleep(10)
url = f"https://{host_ip}/opcua/import_dh_tags"

payload = ""
headers = {
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("PUT", url, headers=headers, data=payload, verify=False)
if response.status_code == 204:
    print("DH tags imported into Hierarchy")

time.sleep(10)
# start the server #
url = f"https://{host_ip}/opcua/start"

payload = {}
headers = {
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)
#print(response)
if response.status_code == 204:
    print("OPC UA server started successfully")

#   Add in the code to use UAExpert or any SSH tools to add OPC UA nodes

#   mount external cifs drive #

url=f'https://{host_ip}/dm/storage/cifs'

payload={
  "name": "test",
  "share": "//10.17.3.4/em63share",
  "readOnly": 'true',
  "mountOnBoot": 'true',
  "username": "litmus1",
  "password": "litmusloop#1"
}

headers = {
    "Authorization": generate_basic_auth_token(api_username_token, api_password_token)
}

response = requests.request("POST", url, headers=headers, data=payload, verify=False)
#print(response)
if response.status_code == 204:
    print("CIFS external driver setup successful")

###      Report word document generation     ###

# Create a new Word document
doc = Document()

# Set the title of the document
doc.add_heading("Litmus Edge Installed on IIoT Gateways", level=1)

doc.add_paragraph("")

# Create a header for the first page
header = doc.sections[0].header
header.paragraphs[0].add_run("Gateway validation Test Report").bold = True
header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# Create the content for the first page
doc.add_heading("Litmus Edge Version:", level=2)
doc.add_heading("Test Report Date:", level=2)

# Add 5 lines of space (blank paragraphs)
for _ in range(2):
    doc.add_paragraph("")

# Create a table with 2 columns
table = doc.add_table(rows=1, cols=2)
table.autofit = False  # Disable autofit to set column widths manually
table.columns[0].width = Pt(200)  # Set the width of the first column
table.columns[1].width = Pt(200)  # Set the width of the second column

# Add a header row to the table
table.rows[0].cells[0].paragraphs[0].add_run("Gateway Details").bold = True
table.rows[0].cells[1].paragraphs[0].add_run("").bold = True

# Define the row entries
entries = [
    "Make:",
    "Model:",
    "Processor:",
    "Memory:",
    "Storage:",
    "Video Interfaces:",
    "Network Interfaces:",
    "USB Ports:",
    "Serial Ports:",
    "GPIO:",
    "Wi-Fi:",
    "BIOS:",
    "Secure Boot:",
    "Power Supply:"
]

# Add rows to the table with the specified entries and empty cells in the second column
for entry in entries:
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(entry)
    row[1].paragraphs[0].add_run("")

# Add a section break to start a new page for the content
doc.add_page_break()

doc.add_paragraph("")

# Create the second page content
doc.add_heading("Contents", level=1)

doc.add_paragraph("")

# Create a table for the Table of Contents
toc_table = doc.add_table(rows=1, cols=2)
toc_table.autofit = False
toc_table.columns[0].width = Pt(350)
toc_table.columns[1].width = Pt(100)

# Add a header row to the Table of Contents
toc_header_cells = toc_table.rows[0].cells
toc_header_cells[0].paragraphs[0].add_run("Title").bold = True
toc_header_cells[1].paragraphs[0].add_run("Page Number").bold = True

# Define the entries for the Table of Contents
toc_entries = [
    "Document Revision History",
    "1.0 Introduction",
    "1.1 Overview",
    "2.0 Tests to be performed",
    "2.1 Devicehub",
    "2.2 Datahub",
    "2.3 Analytics",
    "2.4 Flows Manager",
    "2.5 Integrations",
    "2.6 OPC UA Server",
    "2.7 External Drives",
    "2.8 Perform Network configuration tests",
    "2.9 Wireless capability tests",
    "2.10 Serial Devices",
    "3.0 Additional tests",
    "4.0 Summary of test results",
    "5.0 Major findings during the test",
    "6.0 Recommendations",
]

# Add rows to the Table of Contents with entries and empty page numbers
for entry in toc_entries:
    toc_row = toc_table.add_row().cells
    toc_row[0].paragraphs[0].add_run(entry)
    toc_row[1].paragraphs[0].add_run("")

doc.add_paragraph("")

# Create the "Document Revision History" table

revision_table = doc.add_table(rows=1, cols=4)
revision_table.autofit = False
revision_table.columns[0].width = Pt(80)
revision_table.columns[1].width = Pt(100)
revision_table.columns[2].width = Pt(100)
revision_table.columns[3].width = Pt(270)

# Add a header row to the "Document Revision History" table
revision_header_cells = revision_table.rows[0].cells
revision_header_cells[0].paragraphs[0].add_run("Version #").bold = True
revision_header_cells[1].paragraphs[0].add_run("Revision Date").bold = True
revision_header_cells[2].paragraphs[0].add_run("Revised By").bold = True
revision_header_cells[3].paragraphs[0].add_run("Summary of Changes").bold = True

# Add an empty row to the "Document Revision History" table
revision_table.add_row().cells

#main content (page 3 onwards):

# Add a section break to start a new page for the content
doc.add_page_break()

# Function to add a numbered heading
def add_numbered_heading(heading_number, heading_text):
    p = doc.add_paragraph()
    run = p.add_run(heading_number + " " + heading_text)
    font = run.font
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 255)  # Blue color for headings

# Function to add content for a section
def add_section_content(heading, content):
    add_numbered_heading(heading[0], heading[1:])
    doc.add_paragraph(content)

# Add the content sections
add_section_content("1.0 Introduction", "")
add_section_content("1.1 Overview", "\nThis document provides guidelines to test Litmus Edge software deployed on physical gateways and gather test results. The report provides a summary of the results of tests performed as outlined within this document.")

# Repeat the pattern for other sections (2.0, 2.1, 2.2, ...)

sections = [
    ("2.0 Tests to be Performed:", "\nProject name:\nSystem name:"),
    ("2.1 Devicehub", "\nGuideline: Add 10 devices with up to 500 tags each. See the performance indicators on Dashboard / System Info, observe the thermal condition of the device, and record the results.\n\nTypes of Devices Used: Simulator\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.2 DataHub", "\nGuideline: For each of the ten devices, enable data store in device configuration. Observe the data polling from tags and storing for each device. Add databases and users in Datahub. Check for abnormality in this process and record the results.\n\nData Store enabled: Yes\nAlias Tags Enabled: Yes\nTest Date:\nTest Results:"),
    ("2.3 Analytics", "\nGuideline: Create and run multiple analytics flows. Ensure use of functions with wider WINDOW SIZE for edge analysis and store the results in tsdata & publish to DeviceHub simultaneously. See the performance indicators on Dashboard / System Info, tables getting generated in tsdata, observe the thermal condition of the device, and record the results.\n\nNo. of Analytics Flows Created: 3\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.4 Flows Manager", "\nGuideline: Create multiple and complex flows. Change the memory limit of the flow to 1024 MB. Combine multiple tag data, perform complex calculations, and publish message to DeviceHub node. Ensure use of Analytics and Data Store Nodes with bulk data insertions (InfluxDB or any Marketplace DB App) in Flows. Observe performance indicators and record the results.\n\nNo. of Flows Created:\nDatabases Used: MySQL, PostgreSQL, and InfluxDB\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.5 Integrations", "\nGuideline: Add 2-3 Connectors, preferably one MQTT and remaining with DB Connectors.\n\nNo. of Integrations Created: 2 Integrations Created\nDatabases Used: MySQL DB Connector\nData Streaming Connectors Used: MQTT Integration (LEM)\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.6 OPC UA Server", "\nGuideline: Add DH Tags as OPCUA Nodes, create user for client to connect, start OPCUA server and test client connection. Observe performance indicators and record the results.\n\nNo. of Nodes Created:\nOPCUA Client used for Testing: Unified Automation UaExpert\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.7 External Drives", "\nGuideline: Mount external drives or FTP server and test data transfer from Litmus Edge to the external device. Observe performance indicators and record the results.\n\nType of External Drives: CIFS\nNumber of External Drives Mounted: 1\nFTP Site used for Testing: Ubuntu 18.04 Machine\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.8 Perform Network Configuration Change", "\nGuideline: Try different scenarios by assigning static IP and DHCP. Observe general behavior, performance indicators, and record the results.\n\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.9 Wireless Capability Test", "\nGuideline: Connect to a Wi-Fi Network. Observe general behavior, functionality of LE, performance indicators, and record the results.\n\nTest Date:\nTest Results:\nScreen Shots:"),
    ("2.10 Serial Devices", "\nGuideline: Connect serial devices to serial port or USB port of gateway. Observe general behavior, functionality of LE, performance indicators, and record the results.\n\nType of Serial Device Connected:\nTest Date:\nTest Results:"),
    ("3.0 Additional Tests", "\nGuideline: As needed, perform any additional tests and record the results.\n\nAdditional Test Name:\nDescription:\nTest Date:\nTest Results:\nScreen Shots:"),
    ("4.0 Summary of Test Results", ""),
    ("5.0 Major Findings during Tests", ""),
    ("6.0 Recommendations", "")
]

for section in sections:
    add_section_content(section[0], section[1])

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        font = run.font
        font.name = "Calibri"
        #font.size = Pt(11)

# Save the document
doc.save(r'gateway_validation_report.docx')
print('Report skeleton created')